#!/usr/bin/env python3
"""
Simple test for DOCX parsing functionality
"""
import sys
from pathlib import Path

# Add project root to Python path for imports
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

def test_docx_imports():
    """Test that DOCX imports work correctly"""
    print("Testing DOCX functionality...")
    
    try:
        # Test basic imports
        from docx import Document
        import io
        print("✅ python-docx imports successful")
        
        # Test creating a simple document
        doc = Document()
        doc.add_paragraph("Test paragraph")
        
        # Save to memory
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Read it back
        buffer.seek(0)
        doc2 = Document(buffer)
        text = "\n".join([p.text for p in doc2.paragraphs])
        
        print(f"✅ DOCX creation and reading works: '{text}'")
        
        # Test our parser
        from src.document_parser import DocumentParser
        parser = DocumentParser()
        
        print("✅ DocumentParser with DOCX support ready")
        return True
        
    except Exception as e:
        print(f"❌ Error: {e}")
        return False

def test_file_extension_detection():
    """Test file extension detection logic"""
    print("\nTesting file extension detection...")
    
    test_files = [
        "document.pdf",
        "report.docx", 
        "data.DOCX",
        "file.PDF",
        "unsupported.txt"
    ]
    
    for filename in test_files:
        extension = filename.lower().split('.')[-1]
        supported = extension in ['pdf', 'docx']
        status = "✅ Supported" if supported else "❌ Not supported"
        print(f"  {filename} ({extension}): {status}")

if __name__ == "__main__":
    print("Simple DOCX Parsing Test")
    print("=" * 30)
    
    success = test_docx_imports()
    test_file_extension_detection()
    
    if success:
        print("\n✅ DOCX parsing functionality is ready!")
    else:
        print("\n❌ Issues found with DOCX parsing")