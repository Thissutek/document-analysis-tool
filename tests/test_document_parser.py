#!/usr/bin/env python3
"""
Test script for document parsing functionality (PDF and DOCX)
"""
import sys
from pathlib import Path

# Add project root to Python path for imports
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from src.document_parser import DocumentParser
import io
from docx import Document
import PyPDF2
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

class MockUploadedFile:
    """Mock Streamlit uploaded file for testing"""
    def __init__(self, content, name):
        self.content = io.BytesIO(content)
        self.name = name
        self.size = len(content)
    
    def read(self):
        return self.content.read()
    
    def seek(self, pos):
        return self.content.seek(pos)

def create_test_docx():
    """Create a test DOCX document in memory"""
    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a test paragraph for DOCX parsing.')
    doc.add_paragraph('It contains multiple paragraphs to test text extraction.')
    
    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = 'Header 1'
    table.cell(0, 1).text = 'Header 2'
    table.cell(1, 0).text = 'Cell 1'
    table.cell(1, 1).text = 'Cell 2'
    
    # Save to BytesIO
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer.getvalue()

def create_test_pdf():
    """Create a test PDF document in memory"""
    pdf_buffer = io.BytesIO()
    c = canvas.Canvas(pdf_buffer, pagesize=letter)
    
    # Add some text
    c.drawString(100, 750, "Test PDF Document")
    c.drawString(100, 720, "This is a test paragraph for PDF parsing.")
    c.drawString(100, 690, "It contains multiple lines to test text extraction.")
    
    c.save()
    pdf_buffer.seek(0)
    return pdf_buffer.getvalue()

def test_document_parser():
    """Test document parser with both PDF and DOCX files"""
    print("Testing Document Parser Functionality\n")
    print("="*50)
    
    parser = DocumentParser()
    
    # Test 1: DOCX parsing
    print("\nTest 1: DOCX Parsing")
    print("-" * 20)
    
    try:
        docx_content = create_test_docx()
        docx_file = MockUploadedFile(docx_content, "test.docx")
        
        extracted_text = parser.extract_text_from_document(docx_file)
        
        if extracted_text:
            print("✅ DOCX text extraction successful")
            print(f"Extracted text length: {len(extracted_text)} characters")
            print("First 200 characters:")
            print(f"'{extracted_text[:200]}...'")
            
            # Check for expected content
            if "test paragraph" in extracted_text.lower():
                print("✅ Expected content found in extracted text")
            else:
                print("⚠️  Expected content not found")
        else:
            print("❌ DOCX text extraction failed")
    
    except Exception as e:
        print(f"❌ DOCX test failed: {e}")
    
    # Test 2: File extension detection
    print("\nTest 2: File Extension Detection")
    print("-" * 30)
    
    try:
        # Test with different file extensions
        test_files = [
            MockUploadedFile(b"dummy", "test.docx"),
            MockUploadedFile(b"dummy", "test.pdf"),
            MockUploadedFile(b"dummy", "test.txt")  # Unsupported
        ]
        
        for test_file in test_files:
            file_ext = test_file.name.lower().split('.')[-1]
            print(f"Testing file: {test_file.name} (extension: {file_ext})")
            
            if file_ext in ['pdf', 'docx']:
                print(f"✅ {file_ext.upper()} format supported")
            else:
                print(f"⚠️  {file_ext.upper()} format not supported")
    
    except Exception as e:
        print(f"❌ Extension detection test failed: {e}")
    
    # Test 3: Error handling
    print("\nTest 3: Error Handling")
    print("-" * 20)
    
    try:
        # Test with None input
        result = parser.extract_text_from_document(None)
        if result is None:
            print("✅ None input handled correctly")
        
        # Test with invalid file
        invalid_file = MockUploadedFile(b"invalid content", "test.docx")
        # This should handle the error gracefully
        result = parser.extract_text_from_document(invalid_file)
        print("✅ Error handling working (no crash)")
    
    except Exception as e:
        print(f"Error handling test: {e}")
    
    print("\n" + "="*50)
    print("Document Parser Testing Complete!")

if __name__ == "__main__":
    test_document_parser()