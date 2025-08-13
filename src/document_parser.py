"""
Document Parser Module
Handles PDF and DOCX text extraction using PyPDF2 and python-docx
"""
import PyPDF2
from docx import Document
from typing import Optional
import streamlit as st


class DocumentParser:
    """Extracts text from PDF and DOCX documents"""
    
    def __init__(self):
        pass
    
    def extract_text_from_document(self, uploaded_file) -> Optional[str]:
        """
        Extract text from uploaded document file (PDF or DOCX)
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            str: Extracted text content or None if extraction fails
        """
        if uploaded_file is None:
            return None
        
        file_extension = uploaded_file.name.lower().split('.')[-1]
        
        if file_extension == 'pdf':
            return self.extract_text_from_pdf(uploaded_file)
        elif file_extension == 'docx':
            return self.extract_text_from_docx(uploaded_file)
        else:
            st.error(f"Unsupported file format: {file_extension}")
            return None
    
    def extract_text_from_pdf(self, uploaded_file) -> Optional[str]:
        """
        Extract text from uploaded PDF file
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            str: Extracted text content or None if extraction fails
        """
        try:
            # Read PDF using PyPDF2
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text_content = ""
            
            # Extract text from all pages
            for page in pdf_reader.pages:
                text_content += page.extract_text() + "\n"
            
            # Basic text cleaning
            text_content = self._clean_text(text_content)
            
            return text_content
            
        except Exception as e:
            st.error(f"Error extracting text from PDF: {str(e)}")
            return None
    
    def extract_text_from_docx(self, uploaded_file) -> Optional[str]:
        """
        Extract text from uploaded DOCX file
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            str: Extracted text content or None if extraction fails
        """
        try:
            # Read DOCX using python-docx
            doc = Document(uploaded_file)
            text_content = ""
            
            # Extract text from all paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Skip empty paragraphs
                    text_content += paragraph.text + "\n"
            
            # Extract text from tables if any
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content += cell.text + " "
                    text_content += "\n"
            
            # Basic text cleaning
            text_content = self._clean_text(text_content)
            
            return text_content
            
        except Exception as e:
            st.error(f"Error extracting text from DOCX: {str(e)}")
            return None
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and format extracted text
        
        Args:
            text: Raw extracted text
            
        Returns:
            str: Cleaned text
        """
        # Remove extra whitespace and normalize line breaks
        text = " ".join(text.split())
        
        # Add back some structure
        text = text.replace(". ", ".\n")
        
        return text